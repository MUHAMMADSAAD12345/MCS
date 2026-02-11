"""Document generator tool — creates PDF, DOCX, XLSX files from content."""

from __future__ import annotations

import logging
import time
import uuid
from pathlib import Path
from typing import Any

from config import settings
from models.enums import ToolName
from models.schemas import ToolResult
from tools.base import BaseTool

logger = logging.getLogger(__name__)


class DocumentGeneratorTool(BaseTool):
    """Generate downloadable documents (PDF, DOCX, XLSX) from structured content."""

    @property
    def name(self) -> str:
        return "doc_create"

    async def run(self, **kwargs: Any) -> ToolResult:
        doc_type = kwargs.get("doc_type", "pdf").lower()
        title = kwargs.get("title", "Generated Document")
        content = kwargs.get("content", "")
        data = kwargs.get("data")  # For XLSX: list of dicts

        start = time.monotonic()

        if not content and not data:
            return ToolResult(
                tool=ToolName.DOC_CREATE, success=False, error="No content provided"
            )

        settings.ensure_dirs()
        file_id = str(uuid.uuid4())

        try:
            if doc_type == "pdf":
                filepath = self._generate_pdf(file_id, title, content)
            elif doc_type == "docx":
                filepath = self._generate_docx(file_id, title, content)
            elif doc_type == "xlsx":
                filepath = self._generate_xlsx(file_id, title, data or [])
            else:
                return ToolResult(
                    tool=ToolName.DOC_CREATE,
                    success=False,
                    error=f"Unsupported format: {doc_type}",
                )

            elapsed = (time.monotonic() - start) * 1000
            return ToolResult(
                tool=ToolName.DOC_CREATE,
                success=True,
                data={
                    "file_id": file_id,
                    "filename": filepath.name,
                    "path": str(filepath),
                    "download_url": f"/api/downloads/{file_id}/{filepath.name}",
                },
                latency_ms=round(elapsed, 1),
            )

        except Exception as e:
            elapsed = (time.monotonic() - start) * 1000
            logger.error("Document generation failed: %s", e)
            return ToolResult(
                tool=ToolName.DOC_CREATE,
                success=False,
                error=str(e),
                latency_ms=round(elapsed, 1),
            )

    # ── Generators ────────────────────────────────────────────────────

    @staticmethod
    def _generate_pdf(file_id: str, title: str, content: str) -> Path:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Title
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, title, ln=True, align="C")
        pdf.ln(10)

        # Body
        pdf.set_font("Helvetica", "", 11)
        for line in content.split("\n"):
            pdf.multi_cell(0, 6, line)
            pdf.ln(2)

        filepath = settings.GENERATED_DIR / f"{file_id}.pdf"
        pdf.output(str(filepath))
        return filepath

    @staticmethod
    def _generate_docx(file_id: str, title: str, content: str) -> Path:
        from docx import Document

        doc = Document()
        doc.add_heading(title, level=1)

        for paragraph in content.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

        filepath = settings.GENERATED_DIR / f"{file_id}.docx"
        doc.save(str(filepath))
        return filepath

    @staticmethod
    def _generate_xlsx(file_id: str, title: str, data: list[dict]) -> Path:
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = title[:31]  # Excel sheet name limit

        if data:
            # Headers
            headers = list(data[0].keys())
            ws.append(headers)
            # Rows
            for row in data:
                ws.append([row.get(h, "") for h in headers])
            # Auto-width
            for col_idx, header in enumerate(headers, 1):
                ws.column_dimensions[
                    ws.cell(row=1, column=col_idx).column_letter
                ].width = max(len(header) + 4, 12)

        filepath = settings.GENERATED_DIR / f"{file_id}.xlsx"
        wb.save(str(filepath))
        return filepath
